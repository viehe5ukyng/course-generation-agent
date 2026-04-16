from __future__ import annotations

from pathlib import Path
import subprocess

import fitz
from docx import Document
from rapidocr_onnxruntime import RapidOCR

from app.core.schemas import SourceChunk, SourceDocument


class DocumentParser:
    def __init__(self) -> None:
        self._ocr_engine: RapidOCR | None = None

    def parse_file(self, path: Path, mime_type: str) -> SourceDocument:
        doc = SourceDocument(filename=path.name, mime_type=mime_type)
        try:
            text = self._extract_text(path, mime_type)
            chunks = [
                SourceChunk(text=chunk.strip(), index=index)
                for index, chunk in enumerate(self._chunk_text(text))
                if chunk.strip()
            ]
            doc.text_chunks = chunks
            doc.extract_status = "parsed"
            doc.metadata = {"path": str(path), "chunks": len(chunks)}
        except Exception as exc:  # noqa: BLE001
            doc.extract_status = "failed"
            doc.metadata = {"path": str(path), "error": str(exc)}
        return doc

    def extract_text(self, path: Path, mime_type: str) -> str:
        return self._extract_text(path, mime_type)

    def _extract_text(self, path: Path, mime_type: str) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"} or mime_type.startswith("text/"):
            return path.read_text(encoding="utf-8")
        if suffix == ".pdf":
            return self._extract_pdf_text(path)
        if suffix == ".docx":
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        if suffix == ".doc":
            return self._extract_legacy_word_text(path)
        if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
            return self._ocr_image(path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def _chunk_text(self, text: str, size: int = 1000) -> list[str]:
        return [text[i : i + size] for i in range(0, len(text), size)] or [text]

    def _extract_pdf_text(self, path: Path) -> str:
        doc = fitz.open(path)
        chunks: list[str] = []
        for page in doc:
            extracted = page.get_text("text").strip()
            if extracted:
                chunks.append(extracted)
                continue
            pix = page.get_pixmap(dpi=180)
            ocr_text = self._ocr_bytes(pix.tobytes("png"))
            if ocr_text:
                chunks.append(ocr_text)
        return "\n".join(chunks)

    def _ocr_image(self, path: Path) -> str:
        return self._ocr_bytes(path.read_bytes())

    def _extract_legacy_word_text(self, path: Path) -> str:
        command = ["/usr/bin/textutil", "-convert", "txt", "-stdout", str(path)]
        completed = subprocess.run(command, check=True, capture_output=True, text=True)
        return completed.stdout

    def _ocr_bytes(self, payload: bytes) -> str:
        engine = self._get_ocr_engine()
        result, _ = engine(payload)
        if not result:
            return ""
        lines: list[str] = []
        for item in result:
            if len(item) < 2:
                continue
            text = item[1]
            if isinstance(text, str) and text.strip():
                lines.append(text.strip())
            elif isinstance(text, tuple) and text[0].strip():
                lines.append(text[0].strip())
        return "\n".join(lines)

    def _get_ocr_engine(self) -> RapidOCR:
        if self._ocr_engine is None:
            self._ocr_engine = RapidOCR()
        return self._ocr_engine
